#!/usr/bin/env python3
"""Generate 地块前期研判报告（快评版）Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "D9E2F3")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("地块前期研判报告\n（快评版）")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("产品说明与报告模板")
    run.font.size = Pt(14)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()
    info_lines = [
        "地块名称：________________________",
        "报告日期：________________________",
        "数据截止日：______________________",
        "编制单位：________________________",
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_page_break()

    # Section 0: Overview
    add_heading(doc, "文档说明", 1)
    add_para(
        doc,
        "快评版定位：在 1～3 天内回答核心问题——「这块地值不值得继续往下投（时间/钱/精力）？」"
        "不做完整可研，不做经济测算，不做专家签字担责。适合拿地前初筛、政府更新项目库初选、领导快速汇报。",
    )

    add_heading(doc, "一、快评版回答的核心问题", 2)
    add_table(
        doc,
        ["序号", "问题", "对应章节"],
        [
            ["1", "这块地在哪、长什么样？", "地块概况"],
            ["2", "周边有什么、缺什么？", "周边扫描"],
            ["3", "自身好不好、有什么问题？", "优劣势速判"],
            ["4", "政策允不允许做、有什么红线？", "规划红线"],
            ["5", "值不值得深入、下一步干什么？", "结论建议"],
        ],
        [1.5, 7, 4],
    )

    add_heading(doc, "二、快评版标准内容（5章）", 1)

    # Chapter 1
    add_heading(doc, "第1章  地块概况（约1页）", 2)
    add_table(
        doc,
        ["内容项", "必填", "说明"],
        [
            ["地块位置", "是", "区划、四至、所属更新片区（如有）"],
            ["用地面积", "是", "大致面积（3D测量或公开数据）"],
            ["用地性质", "是", "控规用途（公开数据）"],
            ["现状用途", "是", "空置/工业/住宅/混合"],
            ["现状建筑概况", "是", "栋数、大致层数、建筑年代感"],
            ["3D现状图", "是", "航拍图+3D漫游链接"],
            ["权属/征拆", "标注待核实", "快评不深入产权调查"],
        ],
        [3.5, 2, 8],
    )
    add_para(doc, "【填写区】", bold=True)
    add_para(doc, "地块位置：")
    add_para(doc, "用地面积：          用地性质：          现状用途：")
    add_para(doc, "现状建筑概况：")
    add_para(doc, "3D漫游链接：")
    add_para(doc, "（此处插入四至示意图、3D现状图）")

    # Chapter 2
    add_heading(doc, "第2章  周边扫描（约1～2页）", 2)
    add_heading(doc, "2.1 交通（必填）", 3)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["最近地铁站", "线路+距离（步行分钟）"],
            ["主干道", "是否临主干道/快速路"],
            ["出入口", "地块现有车行出入口情况"],
        ],
        [4, 10],
    )
    add_para(doc, "最近地铁站：          距离：          步行约：    分钟")
    add_para(doc, "主干道情况：          出入口情况：")

    add_heading(doc, "2.2 配套（有/无/距离）", 3)
    add_table(
        doc,
        ["类别", "有/无", "名称", "距离"],
        [
            ["学校", "", "", ""],
            ["医院", "", "", ""],
            ["商业", "", "", ""],
            ["公园", "", "", ""],
            ["地铁/公交", "", "", ""],
            ["负面设施", "", "高架/变电站/污染源等", ""],
        ],
        [2.5, 2, 4, 3],
    )

    add_heading(doc, "2.3 周边建设态势（必填）", 3)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["周边建筑品质", "老旧/混合/较新"],
            ["周边房价参考", "1～2个参考楼盘+单价区间"],
            ["竞品项目", "周边1～2个在售/待售项目"],
            ["天际线/风貌", "3D场景中态势截图"],
        ],
        [4, 10],
    )
    add_para(doc, "（此处插入周边态势3D截图）")

    # Chapter 3
    add_heading(doc, "第3章  优劣势速判（约1页）", 2)
    add_para(doc, "原则：最多列 3 条优势 + 3 条劣势，不展开论述。", bold=True)

    add_heading(doc, "3.1 优势（Strengths）", 3)
    strengths = [
        "□ 区位好（地铁/CBD/产业带）",
        "□ 景观资源（江/河/公园/山）",
        "□ 配套成熟（学区/商业/医疗齐全）",
        "□ 政策红利（更新重点片区/产业扶持）",
        "□ 地形友好（平整/无需大拆）",
        "□ 现状可利用（有保留价值的老建筑）",
    ]
    add_para(doc, "勾选并简述（最多3条）：")
    for s in strengths:
        add_para(doc, s)
    add_para(doc, "优势摘要：")
    add_para(doc, "1. ")
    add_para(doc, "2. ")
    add_para(doc, "3. ")

    add_heading(doc, "3.2 劣势（Weaknesses）", 3)
    weaknesses = [
        "□ 交通不便（无地铁/公交少）",
        "□ 配套缺失（缺学校/商业/医疗）",
        "□ 现状复杂（多产权/租户多/征拆难）",
        "□ 建筑老旧（危房/需拆除重建）",
        "□ 环境负面（临高架/变电站/污染）",
        "□ 市政不足（管网老化/容量不够）",
    ]
    add_para(doc, "勾选并简述（最多3条）：")
    for w in weaknesses:
        add_para(doc, w)
    add_para(doc, "劣势摘要：")
    add_para(doc, "1. ")
    add_para(doc, "2. ")
    add_para(doc, "3. ")

    add_heading(doc, "3.3 现场可见问题（AI识别）", 3)
    issues = [
        "□ 外立面破损/裂缝",
        "□ 疑似违建/加盖",
        "□ 空地/闲置/堆放",
        "□ 消防通道占用",
        "□ 其他：__________",
    ]
    for i in issues:
        add_para(doc, i)
    add_para(doc, "（此处插入AI标注截图）")

    # Chapter 4
    add_heading(doc, "第4章  规划红线（约半页）", 2)
    add_para(doc, "只列「能不能做」的硬约束，不做容量精算。无数据项标注「待核实」。", bold=True)
    add_table(
        doc,
        ["项目", "内容", "来源"],
        [
            ["用地性质", "", "控规公开数据"],
            ["容积率上限", "", "规划条件"],
            ["建筑限高", "", "规划条件"],
            ["是否在更新片区", "是/否，片区名称", "更新专项规划"],
            ["历史保护", "是否涉及历史建筑/风貌区", "文保名录"],
            ["考古风险", "红/黄/蓝/绿", "文物GIS（浙江考古前置）"],
            ["其他红线", "生态红线/河道蓝线/高压走廊", "一张图数据"],
        ],
        [3.5, 5, 4],
    )

    # Chapter 5
    add_heading(doc, "第5章  结论建议（约半页）", 2)

    add_heading(doc, "5.1 综合评级", 3)
    add_table(
        doc,
        ["评级", "含义", "建议"],
        [
            ["🟢 建议深入", "优势明显、红线清晰、值得投入", "进入深评/可研"],
            ["🟡 谨慎推进", "有机会但有明显风险点", "先核实关键事项再决定"],
            ["🔴 暂不建议", "红线限制大/劣势突出", "搁置或调整方向"],
        ],
        [3, 5, 5],
    )
    add_para(doc, "本次评级：□ 🟢建议深入   □ 🟡谨慎推进   □ 🔴暂不建议")

    add_heading(doc, "5.2 推荐方向（一句话）", 3)
    add_para(doc, "示例：建议以拆改结合模式推进，先期综合整治，重点核实征拆意愿和规划条件。")
    add_para(doc, "推荐方向：")

    add_heading(doc, "5.3 下一步待核实事项", 3)
    todos = [
        "□ 规划条件正式查询（规资局）",
        "□ 产权/征拆情况（不动产登记/街道）",
        "□ 土壤环境（如涉及工业用地）",
        "□ 文物考古前置（文物局）",
        "□ 市场深度调研",
        "□ 经济测算（可研）",
    ]
    for t in todos:
        add_para(doc, t)

    add_heading(doc, "5.4 免责声明", 3)
    disclaimer = (
        "本报告为快评版，数据截止日以封面为准。空间数据由三维重建技术生成，"
        "规划及市场数据来自公开渠道，不构成投资决策或规划审批依据。"
        "重大决策须结合深评、官方批复及现场踏勘。"
    )
    p = add_para(doc, disclaimer, size=10)
    for run in p.runs:
        run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # Deliverables
    add_heading(doc, "三、快评版产出物清单", 1)
    add_table(
        doc,
        ["产出物", "格式", "必须"],
        [
            ["快评报告正文", "PDF，3～5页", "是"],
            ["地块3D现状漫游", "Web链接（手机可开）", "是"],
            ["四至区位示意图", "图片1张", "是"],
            ["周边态势截图", "3D场景标注图1～2张", "是"],
            ["SWOT速判表", "表格1张", "是"],
            ["规划红线表", "表格1张", "是"],
            ["综合评级+建议", "半页", "是"],
        ],
        [4.5, 4.5, 2.5],
    )
    add_para(doc, "不包含：经济测算/利润预测、专家签字、详细市场报告、概念方案/总平面图、征拆摸底/产权调查。", bold=True)

    add_heading(doc, "四、快评版 vs 深评版对比", 1)
    add_table(
        doc,
        ["维度", "快评版", "深评版"],
        [
            ["周期", "1～3天", "2～4周"],
            ["页数", "3～5页", "30～50页"],
            ["价格", "0.3～1万", "5～50万"],
            ["核心问题", "值不值得继续？", "怎么做、赚不赚钱？"],
            ["数据来源", "3D扫描+公开数据", "全渠道+现场踏勘+专家"],
            ["责任边界", "辅助参考，不担责", "可含专家签字"],
            ["客户决策", "是否进入下一阶段", "是否拿地/立项"],
            ["自动化程度", "80%+", "40～50%"],
        ],
        [3, 5.5, 5.5],
    )

    add_heading(doc, "五、快评版工作流程", 1)
    workflow = [
        "Day 0 上午：客户提交地块范围；安排无人机/手机拍摄",
        "Day 0 下午～Day 1：3D重建；AI识别建筑、外立面、违建疑似；叠加POI与控规数据",
        "Day 1～Day 2：AI生成优劣势与SWOT；生成规划红线表；人工审核约30分钟",
        "Day 2～Day 3：输出PDF报告+3D漫游链接并交付",
    ]
    for i, step in enumerate(workflow, 1):
        add_para(doc, f"{i}. {step}")

    add_heading(doc, "六、报告页面结构", 1)
    add_table(
        doc,
        ["页码", "章节", "主要内容"],
        [
            ["封面", "—", "地块名称、报告日期、数据截止日、免责声明"],
            ["P1", "地块概况", "基本信息表、四至示意图、3D现状图（二维码）"],
            ["P2", "周边扫描", "交通+配套表、周边态势截图、竞品/房价参考"],
            ["P3", "优劣势速判", "3优势+3劣势、现场可见问题AI标注截图"],
            ["P4", "规划红线", "用地性质/容积率/限高、更新片区/文保/考古风险"],
            ["P5", "结论建议", "综合评级、推荐方向、待核实清单、免责声明"],
        ],
        [2, 3, 9.5],
    )

    add_heading(doc, "七、定价参考", 1)
    add_table(
        doc,
        ["地块规模", "建议定价", "说明"],
        [
            ["单宗 < 5万㎡", "3,000～5,000元", "手机/地面拍摄为主"],
            ["单宗 5～20万㎡", "5,000～10,000元", "需无人机"],
            ["片区级 > 20万㎡", "1～3万元", "多次飞行拼接"],
            ["政府批量（10宗+）", "2,000～3,000元/宗", "规模效应"],
        ],
        [4, 4, 6.5],
    )
    add_para(doc, "商业化路径：快评低价获客 → 引导升级深评（5～20万）→ 政府年框（更新项目库初选）。")

    add_heading(doc, "八、一句话总结", 1)
    add_para(
        doc,
        "快评版 = 3页纸 + 1个3D链接 + 1个评级。"
        "只回答：在哪、周边怎样、好不好、能不能做、值不值得继续。"
        "不做测算、不签字、不担责——但足够支撑「要不要花20万做深评」的决策。",
        bold=True,
    )

    return doc


if __name__ == "__main__":
    output_path = "/workspace/docs/交付模板/地块前期研判报告（快评版）模板.docx"
    import os
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = build_document()
    doc.save(output_path)
    print(f"Saved: {output_path}")
